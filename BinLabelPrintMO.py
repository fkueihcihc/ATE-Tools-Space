import tkinter
from tkinter import *
from tkinter.filedialog import askdirectory
import time
import os
import win32print
import win32api
import win32con
from docx import Document
from docx.shared import Inches
import qrcode
''' 
************************************************************
 FileName:BinLabelPrintMO.py
 Creator: pengbobo <bobo.peng@bitmain.com>
 Create Time: 2022-05-16
 Description: This script is for  printing FT production Label with manual operation
 CopyRight: Copyright belong to bitmain ATE team, All rights reserved.
 '''
# Revision: V1.0.0
# usage: GUI
# ************************************************************

# ************************************************************

def doc_label_create(path,bin,DEVICE,MARKING,LOT_NO,ERP_CODE,TEST_PROGRAM,QTY):
    dataString = DEVICE + '/' + MARKING + '/' + LOT_NO + '/' + TEST_PROGRAM + '/' + QTY
    # 生成二维码
    qrcodePath = path + '\\' + bin +'_'+ QTY +'pcs_' +'二维码.jpg'  # 保存在本地的图片
    res = qrcode.make(data=dataString)  # 生成二维码
    res.save(qrcodePath)  # 二维码保存成文件
    doc_file = path + '\\' + bin +'_' + QTY +'pcs_' + '标签.docx'
    file = Document()  # doc对象
    file.add_paragraph('DEVICE:' + DEVICE + '       ' + bin)
    file.add_paragraph('MARKING:' + MARKING +'   QTY:' + QTY)
    file.add_paragraph('LOT NO:' + LOT_NO)
    file.add_paragraph('ERP CODE:'+ERP_CODE + '  RoHS HF')
    file.add_paragraph('TEST PROGRAM:' + TEST_PROGRAM)
    p = file.add_paragraph('')
    r = p.add_run("     ")  # 段落再增加小节
    r.add_picture(qrcodePath, width=Inches(0.5))  # 小节插入二维码图片
    # file.add_picture(barcode_path, width=Inches(2.0), height=Inches(0.3)) #插入条形码
    file.save(doc_file)  # 保存才能看到结果
    os.remove(qrcodePath)

    return doc_file

def doc_print_device(file_name, printer_name):
    printaccess = {"DesiredAccess": win32print.PRINTER_ACCESS_USE}
    try:
        handle = win32print.OpenPrinter(printer_name, printaccess)
        win32print.ClosePrinter(handle)
        time.sleep(2)
        handle = win32print.OpenPrinter(printer_name, printaccess)
        properties = win32print.GetPrinter(handle, 2)
        devmode = properties['pDevMode']
        devmode.Duplex = 1 # 1：单面，2：双面
        # devmode.PaperSize = 0  #
        # devmode.PaperLength = 40  # .
        # devmode.PaperWidth = 60
        # devmode.Duplex = win32con.DMDUP_SIMPLEX  # 单面
        devmode.Orientation = win32con.DMORIENT_PORTRAIT #DMORIENT_LANDSCAPE  # DMORIENT_PORTRAIT  # 纵向打印
        devmode.Orientation = 1 #1:纵向，2：横向
        # devmode.Scale = 50
        properties['pDevMode'] = devmode
        # for temp_loop in range(2):
        try:
            win32print.SetPrinter(handle, 2, properties, 0)
        except Exception as ff:
            # log.info(f'setprinters:{ff}')
            win32print.SetDefaultPrinter(printer_name)

        # for file_name in file_names:
        # print("file name:",file_name)
        print((f'{printer_name}打印文件{file_name}'))
        win32api.ShellExecute(0, "print", file_name, None, ".", 0)
        time.sleep(2)
        win32print.ClosePrinter(handle)

    except Exception as first:
        # log.info(f'first:{first}')
        print(f'first:{first}')

class MY_GUI():
    def __init__(self,init_window_name,path):
        self.init_window_name = init_window_name
        self.path = path
        # 选择路径函数
    def selectPath(self):
        path_ = askdirectory()
        self.path.set(path_)

    #设置窗口
    def set_init_window(self):
        self.init_window_name.title("BIN标签手动打印V1.0 ©2022 Bitmain")           #窗口名
        self.init_window_name.geometry('425x220+10+10')                 #438x312为窗口大小，+10 +10 定义窗口弹出时的默认展示位置

        # BIN标签
        self.init_bin_label = Label(self.init_window_name, relief=FLAT,justify =RIGHT,text="BIN名称")
        self.init_bin_label.grid(row=1, column=0,columnspan=1)
        # 输入BIN信息
        self.init_bin_entry = Entry(self.init_window_name, width=16)
        self.init_bin_entry.grid(row=1, column=1, rowspan=1,columnspan=1)

        # DEVICE标签
        self.init_device_label = Label(self.init_window_name, relief=FLAT,justify =RIGHT,text="DEVICE")
        self.init_device_label.grid(row=2, column=0, columnspan=1)
        # 输入DEVICE标签
        self.init_device_entry = Entry(self.init_window_name, width=16)
        self.init_device_entry.grid(row=2, column=1, rowspan=1, columnspan=1)

        # MARKING标签
        self.init_marking_label = Label(self.init_window_name, relief=FLAT,justify =RIGHT,text="MARKING")
        self.init_marking_label.grid(row=3, column=0, columnspan=1)
        # 输入MARKING标签
        self.init_marking_entry = Entry(self.init_window_name, width=16)
        self.init_marking_entry.grid(row=3, column=1, rowspan=1, columnspan=1)

        # LOT NO标签
        self.init_lotid_label = Label(self.init_window_name,relief=FLAT,justify =RIGHT, text=" LOT NO")
        self.init_lotid_label.grid(row=4, column=0, columnspan=1)
        # 输入 LOT NO标签
        self.init_lotid_entry = Entry(self.init_window_name, width=16)
        self.init_lotid_entry.grid(row=4, column=1, rowspan=1, columnspan=1)

        # TEST PROGRAM标签
        self.init_testprog_label = Label(self.init_window_name,relief=FLAT,justify =RIGHT, text="TEST PROGRAM")
        self.init_testprog_label.grid(row=5, column=0, columnspan=1)
        # 输入 TEST PROGRAM标签
        self.init_testporg_entry = Entry(self.init_window_name, width=16)
        self.init_testporg_entry.grid(row=5, column=1, rowspan=1, columnspan=1)

        # ERP CODE标签
        self.init_erpcode_label = Label(self.init_window_name,relief=FLAT,justify =RIGHT, text="ERP CODE")
        self.init_erpcode_label.grid(row=6, column=0, columnspan=1)
        # 输入 ERP CODE标签
        self.init_erpcode_entry = Entry(self.init_window_name, width=16)
        self.init_erpcode_entry.grid(row=6, column=1, rowspan=1, columnspan=1)

        # QTY标签
        self.init_qty_label = Label(self.init_window_name, relief=FLAT,justify = RIGHT,text="QTY数量")
        self.init_qty_label.grid(row=7, column=0, columnspan=1)
        # 输入 QTY标签
        self.init_qty_entry = Entry(self.init_window_name, width=16)
        self.init_qty_entry.grid(row=7, column=1, rowspan=2, columnspan=1)

        # 输入 path标签
        self.init_result_entry = Entry(self.init_window_name, width=24,textvariable= self.path)
        self.init_result_entry.grid(row=9, column=1,rowspan=1, columnspan=1)
        #按钮
        self.init_result_button = Button(self.init_window_name,font = ('宋体',8), text="保存标签", bg="lightgrey",
                                         width=8,command=self.selectPath)
        self.init_result_button.grid(row=9, column=2, columnspan=2, sticky=tkinter.W)

        #打印标签按钮
        self.Print_Label = Button(self.init_window_name,font = ('隶书',20,'bold'),fg = 'white',activeforeground = 'grey',
                                relief=RAISED,justify =CENTER, highlightcolor = 'yellow' ,text="打印标签", bg="Blue",
                                height=1,width=9,command=self.Print_Label)  # 调用内部方法  加()为直接调用
        self.Print_Label.grid(row=4, column=2)

    # 功能函数
    # path,bin,DEVICE,MARKING,LOT_NO,ERP_CODE,TEST_PROGRAM,QTY
    def Print_Label(self):
        path = self.init_result_entry.get()
        bin = self.init_bin_entry.get()
        DEVICE = self.init_device_entry.get()
        MARKING=  self.init_marking_entry.get()
        LOT_NO = self.init_lotid_entry.get()
        TEST_PROGRAM = self.init_testporg_entry.get()
        ERP_CODE=  self.init_erpcode_entry.get()
        QTY = self.init_qty_entry.get()

        if path:
                doc_file = doc_label_create(path, bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM, QTY)
                doc_print_device(doc_file,"Citizen CL-F3404")


# the main function
if __name__ == '__main__':
    if 1:
        init_window = Tk()              #实例化出一个父窗口
        path = StringVar()
        ZMJ_PORTAL = MY_GUI(init_window,path)
        # 设置根窗口默认属性
        ZMJ_PORTAL.set_init_window()
        init_window.mainloop()          #父窗口进入事件循环，可以理解为保持窗口运行，否则界面不展示
