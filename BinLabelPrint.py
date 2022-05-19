import tkinter
from tkinter import *
# from tkinter import filedialog
from tkinter.filedialog import askdirectory
import time
import re
import os
import win32print
import win32api
import win32con
from docx import Document
from docx.shared import Inches
import qrcode
# from pystrich.ean13 import EAN13Encoder

''' 
************************************************************
 FileName:BinLabelPrint.py
 Creator: pengbobo <bobo.peng@bitmain.com>
 Create Time: 2022-05-10
 Description: This script is for quickly printing FT production Label
 CopyRight: Copyright belong to bitmain ATE team, All rights reserved.
 '''
# Revision: V1.0.0
# usage: GUI
# ************************************************************

# ************************************************************
# declare global var
debug = 0
LOG_LINE_NUM = 0
'''
class name: BinInfo
Description: This class define a data structure including (paht,file_list)
'''
class BinInfo:
    __slots__ = ["path", "file_list"]

    def __init__(self, p, f):
        self.path, self.file_list = p, f

"""
Function Name: takeFirst()
Description: This function is to sort by first element
"""
def takeFirst(elem):
    return elem.path
"""
Function Name: get_binsum_file()
Description: This function is to get the .sum filename input
"""
def get_binsum_file(path,debug):
    # declare variables
    BI_list = []
    for root, dirs, files in os.walk(path):
        path = root
        sum_file = []
        for f in files:
            if f.endswith('.sum'):
                sum_file.append(f)
        if [] != sum_file:
            sum_file.sort()
            BI = BinInfo(path, sum_file)
            BI_list.append(BI)
    BI_list.sort(key=takeFirst)
    if debug:
        for bi in BI_list:
            print('path:', bi.path)
            print('files:', bi.file_list)

    return BI_list
"""
Function Name: parse_1rst_binsum()
Description: This function is to parse the first test sum file
including test time,program,lot and so on
"""
def parse_1rst_binsum(sum_file, debug):
    flag_sb = 0  # softbin judge
    flag_hb = 0  # hardbin judge
    SBins = {}
    HBins = {'BIN1': '0', 'BIN2': '0', 'BIN3': '0', 'BIN4': '0', 'BIN5': '0',
             'BIN6': '0', 'BIN7': '0', 'BIN8': '0', 'BIN9': '0'}
    tstInfo = []
    f = open(sum_file, 'r')  # read the sum file
    for line in f:
        line = line.strip('\n')
        line = line.strip()
        if not len(line):
            continue
        if line.startswith('Lot: '): #lot id
            test_info = re.match('^Lot:\s*(\S*)', line)
            lot = test_info.group(1)
            tstInfo.append(test_info.group(1))
            if debug: print('lot:', lot)
            continue
        if line.startswith('Started at:'): #test time
            test_info = re.match('.*([0-9]{4}\-[0-9]{2}\-[0-9]{2}).*', line)
            test_time = test_info.group(1)
            tstInfo.append(test_info.group(1))
            if debug: print('test time:', test_time)
            continue
        if line.startswith('Program: '):
            test_info = re.search('^Program:\s+(\w+)', line)
            program = test_info.group(1)
            tstInfo.append(test_info.group(1))
            if debug: print('program:', program)
            continue
        if line.startswith('Tester: '):
            test_info = re.search('^Tester:\s+(\w+)', line)
            tester = test_info.group(1)
            tstInfo.append(test_info.group(1))
            if debug: print('tester:', tester)
            continue
        if line.startswith('CUST: '): # Marking
            test_info = re.search('^CUST:\s+(\w*)', line)
            cust = test_info.group(1)
            tstInfo.append(test_info.group(1))
            if debug: print('cust:', cust)
            continue
        if line.startswith('SUB: '): # sub
            test_info = re.match('^SUB:\s*(\S*)', line)
            sub = test_info.group(1)
            tstInfo.append(test_info.group(1))
            if debug: print('sub:', sub)
            continue
        if line.startswith('TCCT_RDY: '): #ERPCODE
            test_info = re.match('^TCCT_RDY:\s*(\S*)', line)
            erpcode = test_info.group(1)
            tstInfo.append(test_info.group(1))
            if debug: print('erpcode:', erpcode)
            continue
        if line.startswith('Total Number of devices'):
            test_info = re.search('(\d+)', line)
            total_num = test_info.group(1)
            tstInfo.append(test_info.group(1))
            flag_sb = 0
            flag_hb = 0
            if [] == re.findall('B[1-5]C[1-2]', sum_file):
                global base_num
                base_num = int(total_num)
                if debug: print('base_num:', base_num)
            if debug: print('total_num:', total_num)
            continue
        if re.findall('SOFTWARE  BIN  TOTALS', line):
            flag_sb = 1
            flag_hb = 0
            continue
        if re.findall('HARDWARE  BIN  TOTALS', line):
            flag_hb = 1
            flag_sb = 0
            continue
        if 1 == flag_sb and 0 == flag_hb:
            if re.search('\d+', line):
                sb_list = line.split()
                SBins.setdefault('BIN' + sb_list[0], '')
                for i in range(2, len(sb_list) - 1):
                    SBins['BIN' + sb_list[0]] = (sb_list[i])
            continue
        if 1 == flag_hb and 0 == flag_sb:
            if re.search('\d+', line):
                hb_list = line.split()
                HBins['BIN' + hb_list[0]] = hb_list[-2]
            continue

    return tstInfo, dict(sorted(HBins.items()))
"""
Function Name: parse_retest_binsum()
Description: This function is to parse the retest sum file
to get the bin info
"""

def parse_retest_binsum(sum_file, debug):
    flag_sb = 0  # softbin judge
    flag_hb = 0  # hardbin judge
    SBins = {}
    HBins = {'BIN1': '0', 'BIN2': '0', 'BIN3': '0', 'BIN4': '0', 'BIN5': '0',
             'BIN6': '0', 'BIN7': '0', 'BIN8': '0', 'BIN9': '0'}
    f = open(sum_file, 'r')  # read the sum file
    for line in f:
        line = line.strip('\n')
        line = line.strip()
        if not len(line):  # jump space line
            continue
        if re.findall('SOFTWARE  BIN  TOTALS', line):
            flag_sb = 1
            flag_hb = 0
            continue
        if re.findall('HARDWARE  BIN  TOTALS', line):
            flag_hb = 1
            flag_sb = 0
            continue
        if line.startswith('Total Number of devices'):
            flag_sb = 0
            flag_hb = 0
            continue
        if 1 == flag_sb and 0 == flag_hb:
            if re.search('\d+', line):
                sb_list = line.split()
                SBins.setdefault('BIN' + sb_list[0], '')
                for i in range(2, len(sb_list) - 1):
                    SBins['BIN' + sb_list[0]] = (sb_list[i])
            continue
        if 1 == flag_hb and 0 == flag_sb:
            if re.search('\d+', line):
                hb_list = line.split()
                HBins['BIN' + hb_list[0]] = hb_list[-2]
            continue

    return dict(sorted(HBins.items()))
"""
Function Name: Bin_Correct()
Description: This function is to correct the OS BIN number
"""
def Bin_Correct(test_result,OS_BIN):
    # bin calculate
    total_num = 0
    for i in range(1, 9):
        total_num += int(test_result['BIN' + str(i)])
    if total_num > int(test_result['Total number']):
        os_sub = total_num - int(test_result['Total number'])
        if int(test_result[OS_BIN]) - os_sub < 0:
            test_result[OS_BIN] = '0'
        else:
            test_result[OS_BIN] = str(int(test_result[OS_BIN]) - os_sub)

    return test_result

"""
Function Name: check_if_open()
Description: This function is to check if file is open before write into it
"""
def check_if_open(file):
    try:
        os.rename(file, file)
        if debug: print("the file %s is not used,you can use it!!!" % file)
        return True
    except:
        print("the file %s is open,please close it!!!" % file)
        return False

def doc_label_create(path,bin,DEVICE,MARKING,LOT_NO,ERP_CODE,TEST_PROGRAM,QTY):
    dataString = DEVICE + '/' + MARKING + '/' + LOT_NO + '/' + TEST_PROGRAM + '/' + QTY
    # print("dataString:", dataString)
    # # 生成条形码
    # encode = EAN13Encoder('0123456789012')
    # encode.height = Inches(0.3)
    # encode.width = Inches(2.0)
    # barcode_path = path + '\\' + bin + '条形码.png'
    # encode.save(barcode_path)

    # 生成二维码

    qrcodePath = path + '\\' + bin +'_'+ QTY +'pcs_' +'二维码.jpg'  # 保存在本地的图片
    res = qrcode.make(data=dataString)  # 生成二维码
    res.save(qrcodePath)  # 二维码保存成文件
    # print("qrcodePath:",qrcodePath)
    # check if doc file is used
    doc_file = path + '\\' + bin +'_' + QTY +'pcs_' + '标签.docx'
    # if check_if_open(doc_file):
    #     pass
    # else:
    #     doc_file =  path+'\\'+bin+'标签(1).docx'

    file = Document()  # doc对象
    file.add_paragraph('DEVICE:' + DEVICE + '       ' + bin)
    file.add_paragraph('MARKING:' + MARKING +'   QTY:' + QTY)
    file.add_paragraph('LOT NO:' + LOT_NO)
    file.add_paragraph('ERP CODE:'+ERP_CODE + '  RoHS HF')
    file.add_paragraph('TEST PROGRAM:' + TEST_PROGRAM)
    p = file.add_paragraph('')
    r = p.add_run("     ")  # 段落再增加小节
    r.add_picture(qrcodePath, width=Inches(0.5))  # 小节插入二维码图片
    # file.add_picture(barcode_path, width=Inches(2.0), height=Inches(0.3)) #插入条形码
    file.save(doc_file)  # 保存才能看到结果
    os.remove(qrcodePath)

    return doc_file

def doc_print_device(file_name, printer_name):
    printaccess = {"DesiredAccess": win32print.PRINTER_ACCESS_USE}
    try:
        handle = win32print.OpenPrinter(printer_name, printaccess)
        win32print.ClosePrinter(handle)
        time.sleep(2)
        handle = win32print.OpenPrinter(printer_name, printaccess)
        properties = win32print.GetPrinter(handle, 2)
        devmode = properties['pDevMode']
        devmode.Duplex = 1 # 1：单面，2：双面
        devmode.PaperSize = 0  #
        devmode.PaperLength = 40  # .
        devmode.PaperWidth = 60
        # devmode.Duplex = win32con.DMDUP_SIMPLEX  # 单面
        devmode.Orientation = win32con.DMORIENT_PORTRAIT #DMORIENT_LANDSCAPE  # DMORIENT_PORTRAIT  # 纵向打印
        devmode.Orientation = 1 #1:纵向，2：横向
        devmode.Scale = 50
        properties['pDevMode'] = devmode
        # for temp_loop in range(2):
        try:
            win32print.SetPrinter(handle, 2, properties, 0)
        except Exception as ff:
            # log.info(f'setprinters:{ff}')
            win32print.SetDefaultPrinter(printer_name)

        # for file_name in file_names:
        print("file name:",file_name)
        print((f'{printer_name}打印文件{file_name}'))
        win32api.ShellExecute(0, "print", file_name, None, ".", 0)
        time.sleep(2)
        win32print.ClosePrinter(handle)

    except Exception as first:
        # log.info(f'first:{first}')
        print(f'first:{first}')
"""
Function Name: print_label()
Description: This function is to calculate the bin pass/fail number
"""
def print_label(printer_name,test_result,PASS_BINS,path):

    DEVICE = test_result['Mark']
    MARKING = test_result['Sub']
    LOT_NO = test_result['Lot']
    ERP_CODE = test_result['Erpcode']
    TEST_PROGRAM = test_result['Program']

    total_num = int(test_result['Total number'])
    passbin_num = 0

    #case1：print passbin label
    for bin in PASS_BINS:
        QTY = test_result[bin]
        passbin_num +=int(QTY) #calculate passbin number

        #make sure the QTY>0
        if(int(QTY)>0):
            # one tray's max loading number is 3000pcs
            if(int(QTY)<3000): #only one tray
                # step 1 :create file
                doc_file = doc_label_create(path,bin,DEVICE,MARKING,LOT_NO,ERP_CODE,TEST_PROGRAM,QTY)
                # step 2 :print file
                doc_print_device(doc_file,printer_name)

            else: # more than one tray
                if(0 == (int(QTY)%3000) ):
                    # step 1 :create file
                    doc_file = doc_label_create(path,bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM,'3000')
                    # step 2 :print file
                    num_tray = int(QTY) // 3000
                    for i in range(num_tray):
                        # pass
                        doc_print_device(doc_file,printer_name)
                else:
                    num_tray = int(QTY) // 3000 + 1
                    # step 1 :create file
                    doc_file = doc_label_create(path,bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM, '3000')
                    # step 2 :print file
                    num_tray = int(QTY) // 3000
                    for i in range(num_tray-1):
                        # pass
                        doc_print_device(doc_file,printer_name)
                    # for the rest chips
                    # step 1 :create file
                    rest_num = int(QTY)%3000
                    qty_rest = str(rest_num)
                    doc_file = doc_label_create(path,bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM, qty_rest)
                    # step 2 :print file
                    doc_print_device(doc_file,printer_name)

    #case2: print failbin label
    failbin_num = total_num - passbin_num
    QTY = str(failbin_num)
    bin = 'FAILBIN'

    # make sure the QTY>0
    if (int(QTY) > 0):
        # one tray's max loading number is 3000pcs
        if (int(QTY) < 3000):  # only one tray
            # step 1 :create file
            doc_file = doc_label_create(path,bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM, QTY)
            # step 2 :print file
            doc_print_device(doc_file,printer_name)
        else:  # more than one tray
            if (0 == (int(QTY) % 3000)):
                # step 1 :create file
                doc_file = doc_label_create(path,bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM, '3000')
                # step 2 :print file
                num_tray = int(QTY) // 3000
                for i in range(num_tray):
                    # pass
                    doc_print_device(doc_file,printer_name)
            else:
                num_tray = int(QTY) // 3000 + 1
                # step 1 :create file
                doc_file = doc_label_create(path,bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM, '3000')
                # step 2 :print file
                num_tray = int(QTY) // 3000
                for i in range(num_tray - 1):
                    # pass
                    doc_print_device(doc_file,printer_name)
                # for the rest chips
                # step 1 :create file
                rest_num = int(QTY) % 3000
                qty_rest = str(rest_num)
                doc_file = doc_label_create(path,bin, DEVICE, MARKING, LOT_NO, ERP_CODE, TEST_PROGRAM, qty_rest)
                # step 2 :print file
                doc_print_device(doc_file,printer_name)

class MY_GUI():
    def __init__(self,init_window_name,path):
        self.init_window_name = init_window_name
        self.path = path

        # 选择路径函数
    def selectPath(self):
        path_ = askdirectory()
        self.path.set(path_)

    #设置窗口
    def set_init_window(self):
        self.init_window_name.title("BIN标签自动打印V1.0 ©2022 Bitmain")           #窗口名
        self.init_window_name.geometry('438x312+10+10')                 #438x312为窗口大小，+10 +10 定义窗口弹出时的默认展示位置
        # self.init_window_name.geometry('1024x768+10+10')
        # self.init_window_name["bg"] = "Snow"
        # self.init_window_name.attributes("-alpha",0.5)           #虚化，值越小虚化程度越高
        #标签
        # self.log_label = Label(self.init_window_name, text="",font = ('宋体',12))
        # self.log_label.grid(row=2, column=2)

        #导入目录标签
        # self.init_data_label = Label(self.init_window_name, text="选择目录")
        # self.init_data_label.grid(row=0, column=0)
        # 导入目录
        self.init_data_entry = Entry(self.init_window_name, width=16,textvariable= self.path)
        self.init_data_entry.grid(row=1, column=0,rowspan=1, columnspan=1)
        #选择按钮
        self.init_result_button = Button(self.init_window_name,font = ('宋体',10), text="选择LOG文件夹", bg="lightgrey", width=12,command=self.selectPath)
        self.init_result_button.grid(row=1,column=1,rowspan=1,columnspan=1,sticky=tkinter.W)

        #PASS BIN标签
        self.init_passbin_label = Label(self.init_window_name, text="PASS BIN编号")
        self.init_passbin_label.grid(row=2, column=1,columnspan=1)
        # 输入BIN信息
        self.init_passbin_entry = Entry(self.init_window_name, width=16)
        self.init_passbin_entry.grid(row=2, column=0, rowspan=1,columnspan=1)

        # OS BIN标签
        self.init_osbin_label = Label(self.init_window_name, text="OS BIN编号")
        self.init_osbin_label.grid(row=3, column=1, columnspan=1)
        # 输入BIN信息
        self.init_osbin_entry = Entry(self.init_window_name, width=16)
        self.init_osbin_entry.grid(row=3, column=0, rowspan=1, columnspan=1)

        #打印标签按钮
        self.autoPrint = Button(self.init_window_name,font = ('隶书',18,'bold'),fg = 'white',activeforeground = 'grey',
                                relief=RAISED,justify =CENTER, highlightcolor = 'yellow' ,text="打印标签", bg="Blue",
                                height=2,width=8,command=self.autoPrint)  # 调用内部方法  加()为直接调用
        self.autoPrint.grid(row=2, column=5)

        # 输出日志信息框
        self.log_data_Text = Text(self.init_window_name, width=63, height=20, background="white", fg="black")  # 日志框
        self.log_data_Text.grid(row=13, column=0, columnspan=12,rowspan=13)
        self.log_error_Text = Text(self.init_window_name, width=63, height=5, background="white", fg="red")  # 报错日志框
        self.log_error_Text.grid(row=20, column=0, columnspan=12,rowspan=5)

    # 功能函数
    def autoPrint(self):
        global debug, LOG_LINE_NUM
        inputdir = self.init_data_entry.get()
        pass_bins = self.init_passbin_entry.get()
        os_bin =  self.init_osbin_entry.get()
        # print(pass_bins,os_bin)
        if inputdir:
            try:
                self.write_log_to_Text("导入目录："+inputdir)
                BI_list = get_binsum_file(inputdir, debug)
                test_result = {}
                PASS_BINS = []
                test_head = ['Test time', 'Tester','Erpcode','Program', 'Mark', 'Lot', 'Sub', 'Total number']
                if pass_bins:
                    BINS = pass_bins.split(',')
                    for bin in BINS:
                        PASS_BINS.append('BIN'+bin)
                else: # default PASS BIN set
                    PASS_BINS = ['BIN1', 'BIN2', 'BIN3', 'BIN4']
                # print(PASS_BINS)
                if os_bin:
                    OS_BIN = 'BIN'+ os_bin
                else:#default OS BIN set
                    OS_BIN = 'BIN8'
                # print(OS_BIN)
                for bi in BI_list:
                    path = bi.path
                    if debug: print('path:', bi.path)
                    if debug: print('file number:', len(bi.file_list))
                    for sum_file in bi.file_list:
                        # first test file
                        if 1:  print("Input file:", sum_file)
                        sum_file = path + '\\' + sum_file
                        if '-FT-' in sum_file:
                            tstInfo, HBins = parse_1rst_binsum(sum_file, debug)
                            # get the test head info
                            for i in test_head:
                                test_result.setdefault(i, '')
                            for i, j in zip([0,1,2,3,4,5,6,7], [1,3,6,2,4,0,5,7]):  # mapping the information
                                test_result[test_head[i]] = (tstInfo[j])
                            test_result.update(HBins)

                        elif '-RT' in sum_file:
                            HBins = parse_retest_binsum(sum_file, debug)
                            # merge hardbin
                            if '-RT-' in sum_file:
                                for hb in HBins.keys():
                                    if hb in PASS_BINS:
                                        tmp_num = int(test_result[hb]) + int(HBins[hb])
                                        test_result[hb] = str(tmp_num)
                                    else:
                                        tmp_num = int(HBins[hb])
                                        test_result[hb] = str(tmp_num)
                            else:  # more than 1 retest times
                                for hb in HBins.keys():
                                    if OS_BIN == hb:  # OS BIN：replace
                                        tmp_num = int(HBins[hb])
                                        test_result[hb] = str(tmp_num)
                                    else:  # None OS BIN: accumulate
                                        tmp_num = int(float(test_result[hb])) + int(float(HBins[hb]))
                                        test_result[hb] = str(tmp_num)
                    # bin correlation
                    tmp_result = Bin_Correct(test_result,OS_BIN)
                    test_result.update(tmp_result)
                    # print label
                    print_label("Citizen CL-F3404",test_result, PASS_BINS, path)
                LOG_LINE_NUM = 0

                # process finished
                self.write_log_to_Text(">>>>>>打印完成!")

            except:
                self.write_error_to_Text("错误:请先连接打印机或将打印机设为默认打印机！！！")
        else:
            self.write_error_to_Text("错误:请选择导入目录！！！")
    #获取当前时间
    def get_current_time(self):
        current_time = time.strftime('%Y-%m-%d %H:%M:%S',time.localtime(time.time()))
        return current_time

    #日志动态打印
    def write_log_to_Text(self,logmsg):
        global LOG_LINE_NUM
        current_time = self.get_current_time()
        logmsg_in = str(current_time) +" " + str(logmsg) + "\n"      #换行
        if LOG_LINE_NUM <= 120:
            self.log_data_Text.insert(END, logmsg_in)
            LOG_LINE_NUM = LOG_LINE_NUM + 1
        else:
            self.log_data_Text.delete(1.0,2.0)
            self.log_data_Text.insert(END, logmsg_in)

    #日志动态打印
    def write_error_to_Text(self,logmsg):
        global LOG_LINE_NUM
        current_time = self.get_current_time()
        logmsg_in = str(current_time) +" " + str(logmsg) + "\n"      #换行
        if LOG_LINE_NUM <= 120:
            self.log_error_Text.insert(END, logmsg_in)
            LOG_LINE_NUM = LOG_LINE_NUM + 1
        else:
            self.log_error_Text.delete(1.0,2.0)
            self.log_error_Text.insert(END, logmsg_in)

# the main function
if __name__ == '__main__':
    if 1:
        init_window = Tk()              #实例化出一个父窗口
        path = StringVar()
        ZMJ_PORTAL = MY_GUI(init_window,path)
        # 设置根窗口默认属性
        ZMJ_PORTAL.set_init_window()
        init_window.mainloop()          #父窗口进入事件循环，可以理解为保持窗口运行，否则界面不展示
    if 0:
        inputdir = r'D:\01_myWork\02_ScriptsDev\BinLabelPrint\BM1764AB\20220224'
        BI_list = get_binsum_file(inputdir, debug)
        test_result = {}
        test_head = ['Test time', 'Tester','Erpcode','Program', 'Mark', 'Lot', 'Sub', 'Total number']
        PASS_BINS = ['BIN1', 'BIN2', 'BIN3', 'BIN4']
        OS_BIN = 'BIN8'
        for bi in BI_list:
            path = bi.path
            if debug: print('path:', bi.path)
            if debug: print('file number:', len(bi.file_list))
            for sum_file in bi.file_list:
                # first test file
                if 1:  print("Input file:", sum_file)
                sum_file = path + '\\' + sum_file
                if '-FT-' in sum_file:
                    tstInfo, HBins = parse_1rst_binsum(sum_file, debug)
                    # get the test head info
                    for i in test_head:
                        test_result.setdefault(i, '')
                    for i, j in zip([0,1,2,3,4,5,6,7], [1,3,6,2,4,0,5,7]):  # mapping the information
                        test_result[test_head[i]] = (tstInfo[j])
                    test_result.update(HBins)

                elif '-RT' in sum_file:
                    HBins = parse_retest_binsum(sum_file, debug)
                    # merge hardbin
                    if '-RT-' in sum_file:
                        for hb in HBins.keys():
                            if hb in PASS_BINS:
                                tmp_num = int(test_result[hb]) + int(HBins[hb])
                                test_result[hb] = str(tmp_num)
                            else:
                                tmp_num = int(HBins[hb])
                                test_result[hb] = str(tmp_num)
                    else:  # more than 1 retest times
                        for hb in HBins.keys():
                            if OS_BIN == hb:  # OS BIN：replace
                                tmp_num = int(HBins[hb])
                                test_result[hb] = str(tmp_num)
                            else:  # None OS BIN: accumulate
                                tmp_num = int(float(test_result[hb])) + int(float(HBins[hb]))
                                test_result[hb] = str(tmp_num)
            # bin correlation
            tmp_result = Bin_Correct(test_result,OS_BIN)
            test_result.update(tmp_result)
            # print label
            print_label("Citizen CL-F3404", test_result, PASS_BINS, path)